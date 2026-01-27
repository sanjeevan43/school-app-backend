
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from fpdf import FPDF
import os

def clean_text_for_pdf(text):
    # Remove emojis and non-latin characters for simple FPDF usage
    # This is a basic filter to prevent encoding errors with standard fonts
    return text.encode('latin-1', 'ignore').decode('latin-1')

def add_docx_styles(doc):
    # Ensure styles exist
    styles = doc.styles
    # You can customize styles here if needed
    pass

def markdown_to_docx(md_lines, output_file):
    doc = Document()
    add_docx_styles(doc)
    
    in_code_block = False
    code_content = []
    
    for line in md_lines:
        line = line.rstrip()
        
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                p = doc.add_paragraph()
                p.style = 'No Spacing'
                runner = p.add_run('\n'.join(code_content))
                runner.font.name = 'Courier New'
                runner.font.size = Pt(9)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            continue
            
        if in_code_block:
            code_content.append(line)
            continue
        
        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('---') or line.startswith('___'):
             doc.add_paragraph('________________________________________________________________')
        else:
            # Basic Bold processing
            parts = re.split(r'(\*\*.*?\*\*)', line)
            p = doc.add_paragraph()
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    doc.save(output_file)
    print(f"Created {output_file}")

class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 10)
        self.cell(0, 10, 'School Transport API Documentation', 0, 1, 'R')
        self.ln(5)

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, 'Page ' + str(self.page_no()) + '/{nb}', 0, 0, 'C')

def markdown_to_pdf(md_lines, output_file):
    pdf = PDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.set_font("Arial", size=11)
    
    in_code_block = False
    
    for line in md_lines:
        line = line.rstrip()
        clean_line = clean_text_for_pdf(line)
        
        if line.startswith('```'):
            in_code_block = not in_code_block
            if in_code_block:
                pdf.ln(2)
                pdf.set_font("Courier", size=9)
            else:
                pdf.ln(2)
                pdf.set_font("Arial", size=11)
            continue
            
        if in_code_block:
            pdf.multi_cell(0, 5, clean_line)
            continue
            
        if line.startswith('# '):
            pdf.ln(5)
            pdf.set_font("Arial", 'B', 16)
            pdf.multi_cell(0, 10, clean_line[2:])
            pdf.set_font("Arial", size=11)
        elif line.startswith('## '):
            pdf.ln(4)
            pdf.set_font("Arial", 'B', 14)
            pdf.multi_cell(0, 10, clean_line[3:])
            pdf.set_font("Arial", size=11)
        elif line.startswith('### '):
            pdf.ln(3)
            pdf.set_font("Arial", 'B', 12)
            pdf.multi_cell(0, 10, clean_line[4:])
            pdf.set_font("Arial", size=11)
        elif line.startswith('---'):
            pdf.ln(5)
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(5)
        else:
            # Handle Bold regex for PDF is harder, simplified here to just text
            # Splitting by bold would require write(), but multi_cell is safer for wrapping
            # We'll stick to multi_cell for stability
            pdf.multi_cell(0, 6, clean_line)

    pdf.output(output_file)
    print(f"Created {output_file}")

if __name__ == "__main__":
    file_path = 'COMPLETE_API_DOCUMENTATION.md'
    if os.path.exists(file_path):
        with open(file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        try:
            markdown_to_docx(lines, 'School_Transport_API_Docs.docx')
        except Exception as e:
            print(f"Error creating DOCX: {e}")
            
        try:
            markdown_to_pdf(lines, 'School_Transport_API_Docs.pdf')
        except Exception as e:
            print(f"Error creating PDF: {e}")
    else:
        print(f"File {file_path} not found.")
